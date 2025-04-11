import streamlit as st
import cv2
import numpy as np
import os
import tempfile
import base64
import requests
import json
import glob
from PIL import Image
from openai import OpenAI
from docx import Document
from streamlit_drawable_canvas import st_canvas
from dotenv import load_dotenv

# Set Streamlit theme - Must be the first Streamlit command
st.set_page_config(page_title="VT Generator", page_icon="🖼️", layout="wide")
load_dotenv()

# Initialize OpenAI client
GPT_API_KEY = os.getenv("PERSONAL_OPENAI_KEY")
client = OpenAI(api_key=GPT_API_KEY)

# --- Helper Functions ---
# The line below was causing a NameError - removing or properly commenting it

# Define minimal fallback functions if needed
def image_to_base64(image):
    """Convert PIL image or numpy array to base64 string."""
    if isinstance(image, np.ndarray):
        image = Image.fromarray(image)
    buffered = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
    image.save(buffered, format="JPEG")
    with open(buffered.name, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")

def get_frame_timestamp(frame_number, video_obj):
    """Get timestamp for a frame."""
    if video_obj and video_obj.isOpened():
        fps = video_obj.get(cv2.CAP_PROP_FPS)
        if fps > 0:
            seconds = frame_number / fps
            return seconds
    return 0

# Function to encode image as base64
def encode_image(image):
    buffered = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
    image.save(buffered, format="JPEG")
    with open(buffered.name, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")

# Function to parse SRT files
def parse_srt(file):
    subtitles = {}
    lines = file.read().decode("utf-8").split("\n")
    index, start_time = None, None
    for line in lines:
        line = line.strip()
        if line.isdigit():
            index = int(line)
        elif "-->" in line:
            start_time = line.split(" --> ")[0]
            start_time = sum(float(x) * 60 ** i for i, x in enumerate(reversed(start_time.replace(',', '.').split(':'))))
        elif line:
            if index is not None and start_time is not None:
                subtitles[start_time] = line
    return subtitles

# --- Cropping Logic Functions ---
def crop_rectangular(image_cv_bgr, rect_data):
    """Crops the OpenCV BGR image using rectangle data."""
    left = int(rect_data['left'])
    top = int(rect_data['top'])
    width = int(rect_data['width'])
    height = int(rect_data['height'])
    if width <= 0 or height <= 0:
        st.warning("Please draw a valid rectangle.")
        return None
    h_img, w_img = image_cv_bgr.shape[:2]
    x1, y1 = max(0, left), max(0, top)
    x2, y2 = min(w_img, left + width), min(h_img, top + height)
    if x2 <= x1 or y2 <= y1:
         st.warning("Calculated crop area is outside image bounds or invalid.")
         return None
    cropped_bgr = image_cv_bgr[y1:y2, x1:x2]
    return cropped_bgr

def crop_freeform(image_cv_bgr, path_data):
    """Crops the OpenCV BGR image using freeform path data."""
    if not path_data:
         st.warning("Received empty path data.")
         return None
    
    # Get image dimensions for boundary validation
    h_img, w_img = image_cv_bgr.shape[:2]
    
    points_list = []
    for point_cmd in path_data:
        if len(point_cmd) >= 3:
            try:
                # Extract coordinates
                x = int(float(point_cmd[-2]))
                y = int(float(point_cmd[-1]))
                
                # Clip coordinates to image boundaries
                x = max(0, min(x, w_img - 1))
                y = max(0, min(y, h_img - 1))
                
                # Add the valid point to our list
                points_list.append([x, y])
            except (ValueError, IndexError, TypeError):
                # Skip invalid point data silently
                continue

    # Ensure we have enough points to form a shape
    if len(points_list) < 3:
         st.warning("Not enough valid points to create a crop area. Please try again.")
         return None
    
    try:
        # Convert to numpy array for OpenCV
        contour = np.array(points_list, dtype=np.int32)
        
        # Create a mask with only the points inside the image
        mask = np.zeros(image_cv_bgr.shape[:2], dtype=np.uint8)
        cv2.drawContours(mask, [contour], -1, color=255, thickness=cv2.FILLED)
        
        # Apply the mask
        masked_image_bgr = cv2.bitwise_and(image_cv_bgr, image_cv_bgr, mask=mask)
        
        # Calculate bounding rectangle
        x_bb, y_bb, w_bb, h_bb = cv2.boundingRect(contour)
        
        # Validate bounding box dimensions
        if w_bb <= 0 or h_bb <= 0:
            st.warning("Freeform crop area resulted in an empty image. Please try a larger selection.")
            return None
            
        # Ensure bounding box is within image boundaries
        x_bb = max(0, x_bb)
        y_bb = max(0, y_bb)
        w_bb = min(w_bb, w_img - x_bb)
        h_bb = min(h_bb, h_img - y_bb)
        
        # Final validation of crop area
        if w_bb <= 0 or h_bb <= 0:
            st.warning("Crop area is outside image boundaries. Please try again.")
            return None
            
        # Extract the bounded region
        cropped_bgr = masked_image_bgr[y_bb:y_bb+h_bb, x_bb:x_bb+w_bb]
        return cropped_bgr
    
    except Exception as e:
        st.error(f"Error during freeform cropping: {e}")
        return None

# Function to get list of users from the database directory
def get_settings():
    """Load default settings from file if available, otherwise return defaults"""
    settings_path = "visual_transcription/database/default.json"
    try:
        if os.path.exists(settings_path):
            with open(settings_path, 'r') as f:
                settings = json.load(f)
                
            # Load navigation settings
            if 'frame_increment' in settings:
                st.session_state.frame_increment = settings['frame_increment']
                
            # Load drawing settings
            if 'stroke_slider' in settings:
                st.session_state.stroke_slider = settings['stroke_slider']
            if 'stroke_color' in settings:
                st.session_state.stroke_color = settings['stroke_color']
    except Exception as e:
        st.error(f"Error loading settings: {e}")
        # Use defaults if settings can't be loaded

def save_settings():
    """Save current settings to default file"""
    # Create the directory if it doesn't exist
    os.makedirs("visual_transcription/database", exist_ok=True)
    
    settings_path = "visual_transcription/database/default.json"
    try:
        # Create a settings dictionary with just navigation and drawing settings
        settings = {
            'frame_increment': st.session_state.get('frame_increment', 1),
            'stroke_slider': st.session_state.get('stroke_slider', 3),
            'stroke_color': st.session_state.get('stroke_color', '#00FF00')
        }
        
        with open(settings_path, 'w') as f:
            json.dump(settings, f, indent=4)
        
        return True
    except Exception as e:
        st.error(f"Error saving settings: {e}")
        return False

# Download full transcript
def download_transcript():
    doc = Document()
    doc.add_heading("Visual Transcript", level=1)
    
    # Use the merged transcripts for a more comprehensive document
    merged_transcripts = merge_transcripts()
    
    if merged_transcripts:
        # Add a section explaining the format
        doc.add_paragraph("This document contains both audio transcripts and visual descriptions in chronological order.")
        doc.add_paragraph("Timestamps are shown in seconds from the start of the video.")
        
        # Add all merged transcripts in chronological order
        for entry in merged_transcripts:
            if len(entry) == 3:  # Audio entry (timestamp, text, "audio")
                timestamp, text, source = entry
                para = doc.add_paragraph()
                para.add_run(f"[{timestamp:.2f}s] ").bold = True
                para.add_run(f"{text}")
            else:  # Visual entry (timestamp, text, "visual", frame_number)
                timestamp, text, source, frame_number = entry
                para = doc.add_paragraph()
                para.add_run(f"[{timestamp:.2f}s - Frame {frame_number}] ").bold = True
                para.add_run(f"Visual Description: {text}").italic = True
    else:
        # Fall back to original subtitles if no merged transcripts
        for timestamp, text in st.session_state["subtitles"].items():
            doc.add_paragraph(f"{timestamp}: {text}")
    
    temp_doc_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(temp_doc_path)
    with open(temp_doc_path, "rb") as doc_file:
        st.sidebar.download_button("Download Transcript", doc_file, file_name="visual_transcript.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Function to merge audio and visual transcripts chronologically
def merge_transcripts():
    """
    Creates a chronologically ordered combined transcript from both 
    audio subtitles and visual frame transcriptions.
    Returns a list of (timestamp, text, source_type) tuples sorted by timestamp.
    """
    combined_entries = []
    
    # Add entries from audio transcripts (subtitles)
    if st.session_state["subtitles"]:
        for timestamp, text in st.session_state["subtitles"].items():
            # Handle potential format issues
            try:
                combined_entries.append((float(timestamp), text, "audio"))
            except (ValueError, TypeError):
                # If timestamp can't be converted to float, use a default
                st.warning(f"Invalid timestamp format: {timestamp}. Using 0.0 instead.")
                combined_entries.append((0.0, text, "audio"))
    
    # Add entries from visual transcriptions
    if st.session_state.get("transcriptions") and st.session_state.get("inserted_transcriptions"):
        # Get video object for timestamp conversion
        video_obj = st.session_state.get('video')
        if video_obj and video_obj.isOpened():
            for frame_number in st.session_state.inserted_transcriptions:
                if frame_number in st.session_state["transcriptions"]:
                    # Convert frame number to timestamp
                    timestamp = get_frame_timestamp(frame_number, video_obj)
                    transcription = st.session_state["transcriptions"][frame_number]
                    combined_entries.append((timestamp, transcription, "visual", frame_number))
        else:
            # If video isn't available, use placeholder timestamps
            for frame_number in st.session_state.inserted_transcriptions:
                if frame_number in st.session_state["transcriptions"]:
                    # Use frame number as approximate timestamp
                    transcription = st.session_state["transcriptions"][frame_number]
                    combined_entries.append((float(frame_number)/30.0, transcription, "visual", frame_number))
    
    # Sort by timestamp (first element of each tuple)
    return sorted(combined_entries, key=lambda x: x[0])

# Initialize session state variables
# Original session state variables
st.session_state.setdefault("saved_frames", {})
st.session_state.setdefault("saved_subtitles", [])
st.session_state.setdefault("frame_index", 0)
st.session_state.setdefault("frame_subtitle_map", {})
st.session_state.setdefault("subtitles", {})
st.session_state.setdefault("transcriptions", {})
# New session state to track which transcriptions have been inserted into the transcript
st.session_state.setdefault("inserted_transcriptions", set())

# Additional session state for enhanced features
st.session_state.setdefault("video", None)
st.session_state.setdefault("frame_number", 0)
st.session_state.setdefault("total_frames", 0)
st.session_state.setdefault("uploaded", False)
st.session_state.setdefault("audio_transcript", [])
st.session_state.setdefault("canvas_key", 0)
st.session_state.setdefault("frame_increment", 1)
st.session_state.setdefault("show_settings", True)
st.session_state.setdefault("pending_video_file", None)
st.session_state.setdefault("stroke_slider", 3)
st.session_state.setdefault("stroke_color", "#00FF00")
# Add active tab tracking
st.session_state.setdefault("active_tab", 0)  # 0=Settings, 1=Media Upload, 2=Visual Transcription
# Add flag for showing workspace after video processing
st.session_state.setdefault("show_workspace", False)
st.session_state["max_words"] = st.session_state.get("max_words", "20")
# Try to load settings
get_settings()

# GPT-4o settings
if "gpt-4o" not in st.session_state:
    try:
        # First try to load default (general) prompt from new prompt system
        prompt_file_path = "database/prompts/general.json"
        if os.path.exists(prompt_file_path):
            with open(prompt_file_path, "r") as json_file:
                gpt4o_data = json.load(json_file)
                st.session_state['gpt-4o'] = {
                    "prompt": gpt4o_data["prompt"], 
                    "max_words": st.session_state["max_words"]
                }
                # Initial replacement
                st.session_state['gpt-4o']["prompt"] = st.session_state['gpt-4o']["prompt"].replace(
                    "%MAX_WORDS%", str(st.session_state["max_words"]))
        # Fall back to legacy path for backward compatibility
        elif os.path.exists(r"utils\chat_GPT.json"):
            with open(r"utils\chat_GPT.json", "r") as json_file:
                gpt4o_data = json.load(json_file)
                st.session_state['gpt-4o'] = {
                    "prompt": gpt4o_data["prompt"], 
                    "max_words": st.session_state["max_words"]
                }
                # Initial replacement
                st.session_state['gpt-4o']["prompt"] = st.session_state['gpt-4o']["prompt"].replace(
                    "%MAX_WORDS%", str(st.session_state["max_words"]))
        else:
            # Default fallback if no files found
            st.session_state['gpt-4o'] = {"prompt": "Describe the image in detail.", "max_words": "100"}
        
        # Initialize prompt category
        st.session_state.prompt_category = "general"
            
    except Exception as e:
        st.error(f"Error loading prompt settings: {e}")
        st.session_state['gpt-4o'] = {"prompt": "Describe the image in detail.", "max_words": "100"}

# Sidebar setup
st.sidebar.title("Saved Frames & Transcripts")
# The sidebar content will be updated later in the file

# --- App Title and Initial Setup ---
st.title('Video Transcription Service with Cropping')

# Create three main tabs for the entire application
tab_names = ["Configuration Settings", "Media Upload", "Visual Transcription"]
settings_tab, media_tab, workspace_tab = st.tabs(tab_names)

# If we just processed a video, we need to navigate to the workspace tab
# This is done using the show_workspace flag in session state
if st.session_state.show_workspace and st.session_state.uploaded:
    # Reset the flag so we don't keep redirecting
    st.session_state.show_workspace = False
    # Put the Visual Transcription tab content directly here to display first
    st.warning("Video processed! Navigating to Visual Transcription tab...")
    # Using rerun often causes issues with tab selection, so instead we'll
    # create a message instructing the user
    st.success("Please click on the 'Visual Transcription' tab to view your video.")

# -----------------------------------------------
# TAB 1: CONFIGURATION SETTINGS
# -----------------------------------------------
with settings_tab:
    # Navigation Settings
    st.subheader("Navigation Settings")
    # Frame increment input
    increment_value = st.number_input(
        'Frame increment (number of frames to jump when using navigation buttons)',
        min_value=1,
        max_value=100,
        value=st.session_state.frame_increment,
        step=1,
        key='increment_input'
    )
    # Update session state with the new increment value
    if increment_value != st.session_state.frame_increment:
        st.session_state.frame_increment = int(increment_value)
        st.success(f"Frame increment updated to {st.session_state.frame_increment}")
    
    st.markdown("---")
    
    # Drawing Settings
    st.subheader("Drawing Settings")
    # Add drawing settings
    # Add stroke width slider
    stroke_width = st.slider("Stroke width:", 1, 25, st.session_state.stroke_slider, key='stroke_width_input')
    if stroke_width != st.session_state.stroke_slider:
        st.session_state.stroke_slider = stroke_width
        
    # Add color picker for stroke color
    stroke_color = st.color_picker("Stroke color:", st.session_state.stroke_color, key='stroke_color_input')
    if stroke_color != st.session_state.stroke_color:
        st.session_state.stroke_color = stroke_color
    
    st.markdown("---")
    
    # GPT-4o Settings
    st.subheader("GPT-4o Settings")
    # Display and allow editing of the max words
    prev_max_words = str(st.session_state['gpt-4o'].get("max_words", "100"))
    max_words_input = st.text_input(
        label="Max words for description",
        value=prev_max_words,
        key='gpt4_max_words'
    )
    current_max_words = str(max_words_input)
    st.session_state["max_words"] = current_max_words

    # Update prompt if max_words changed
    current_prompt = st.session_state['gpt-4o'].get("prompt", "Describe the image in %MAX_WORDS% words.")
    if prev_max_words != current_max_words:
        st.session_state['gpt-4o']["prompt"] = current_prompt.replace(f"{prev_max_words}", f"{current_max_words}")
        st.session_state['gpt-4o']["max_words"] = current_max_words
        current_prompt = st.session_state['gpt-4o']["prompt"]
    
    # Save settings button
    st.markdown("---")
    if st.button("Save Configuration", key="save_config_button"):
        if save_settings():
            st.success("Settings saved successfully")

# -----------------------------------------------
# TAB 2: MEDIA UPLOAD
# -----------------------------------------------
with media_tab:
    # Audio Transcript
    st.subheader("Audio Transcript")
    # Load Audio Transcript if not already loaded
    if not st.session_state.audio_transcript:
        # Replace the old uploader with the SRT file uploader
        srt_file = st.file_uploader("Upload Subtitle File (SRT)", type=["srt"], key='srt_uploader')
        if srt_file is not None:
            # Process SRT file immediately if video already loaded 
            if st.session_state.get('uploaded', False) and st.session_state.get('video') is not None:
                try:
                    # Save current transcriptions and inserted_transcriptions
                    saved_transcriptions = st.session_state.get("transcriptions", {})
                    saved_inserted_transcriptions = st.session_state.get("inserted_transcriptions", set())
                    
                    # Parse and process the SRT file
                    st.session_state["subtitles"] = parse_srt(srt_file)
                    fps = st.session_state.video.get(cv2.CAP_PROP_FPS)
                    if fps > 0:
                        st.session_state["frame_subtitle_map"] = {
                            int(start_time * fps): text
                            for start_time, text in st.session_state["subtitles"].items()
                        }
                    else:
                        st.warning("Could not get FPS from video. Subtitle mapping might be incorrect.")
                    
                    # Restore saved transcriptions and inserted_transcriptions
                    st.session_state["transcriptions"] = saved_transcriptions
                    st.session_state["inserted_transcriptions"] = saved_inserted_transcriptions
                    
                    st.success("SRT file processed! Visual and audio transcripts will be displayed together.")
                    
                except Exception as e:
                    st.error(f"Error processing SRT file: {e}")
            else:
                st.success('SRT file upload detected! File will be processed when you apply settings.')
    else:
        st.success("✅ Audio transcript loaded")
        if st.button("Clear Audio Transcript"):
            st.session_state.audio_transcript = []
            st.experimental_rerun()
    
    st.markdown("---")
    
    # Video and Subtitle Upload
    st.subheader("Video Upload")
    # Video uploader
    if not st.session_state.uploaded:
        uploaded_file = st.file_uploader('Drag and drop a video file', type=['mp4', 'avi', 'mov'], key='video_uploader')
        if uploaded_file is not None:
            st.session_state.pending_video_file = uploaded_file
            st.success('Video upload detected! Video will be processed when you apply settings.')
    else:
        st.success(f"✅ Video loaded ({st.session_state.total_frames} frames)")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Clear Video"):
                if st.session_state.video is not None:
                    st.session_state.video.release()
                st.session_state.video = None
                st.session_state.uploaded = False
                st.session_state.saved_frames = {}
                st.session_state.frame_number = 0
                st.session_state.total_frames = 0
                st.session_state.pending_video_file = None
                st.session_state.subtitles = {}
                # Keep transcriptions in case user wants to reuse them
                # st.session_state.transcriptions = {}
                # st.session_state.inserted_transcriptions = set()
                st.experimental_rerun()
        
        with col2:
            if st.button("Clear Video & Transcriptions"):
                if st.session_state.video is not None:
                    st.session_state.video.release()
                st.session_state.video = None
                st.session_state.uploaded = False
                st.session_state.saved_frames = {}
                st.session_state.frame_number = 0
                st.session_state.total_frames = 0
                st.session_state.pending_video_file = None
                st.session_state.subtitles = {}
                # Clear all transcriptions
                st.session_state.transcriptions = {}
                st.session_state.inserted_transcriptions = set()
                st.experimental_rerun()
    
    st.markdown("---")
    # Create a 3-column layout with empty space on the left to push button to right
    _, _, right_col = st.columns([3, 2, 2])
    with right_col:
        # Only show Process Media button if a video is uploaded and pending processing
        if st.session_state.pending_video_file is not None:
            if st.button("Process Media", key="process_media_button", type="primary"):
                # Process the pending video file if it exists
                if not st.session_state.uploaded:
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.mp4') as temp_file:
                        temp_file.write(st.session_state.pending_video_file.read())
                        temp_file_path = temp_file.name
                    st.write(f'Video saved temporarily.')
                    if not os.path.exists(temp_file_path):
                        st.error('Temporary file was not created successfully.')
                        st.session_state.uploaded = False
                        st.session_state.video = None
                    else:
                        # Clear previous state before loading new video
                        st.session_state.saved_frames = {}
                        st.session_state.frame_number = 0 # Reset frame number
                        if st.session_state.video is not None:
                            st.session_state.video.release()

                        st.session_state.video = cv2.VideoCapture(temp_file_path)
                        if st.session_state.video.isOpened():
                            st.session_state.total_frames = int(st.session_state.video.get(cv2.CAP_PROP_FRAME_COUNT))
                            st.session_state.uploaded = True

                            # Process SRT file if it exists
                            srt_file = st.session_state.get('srt_uploader')
                            # Save current transcriptions and inserted_transcriptions to restore after SRT processing
                            saved_transcriptions = st.session_state.get("transcriptions", {})
                            saved_inserted_transcriptions = st.session_state.get("inserted_transcriptions", set())
                            
                            st.session_state["subtitles"] = {} # Ensure subtitles dict exists
                            st.session_state["frame_subtitle_map"] = {} # Ensure map exists
                            if srt_file is not None:
                                try:
                                    st.session_state["subtitles"] = parse_srt(srt_file)
                                    fps = st.session_state.video.get(cv2.CAP_PROP_FPS)
                                    if fps > 0:
                                        st.session_state["frame_subtitle_map"] = {
                                            int(start_time * fps): text
                                            for start_time, text in st.session_state["subtitles"].items()
                                        }
                                    else:
                                        st.warning("Could not get FPS from video. Subtitle mapping might be incorrect.")
                                except Exception as e:
                                    st.error(f"Error parsing SRT file: {e}")
                            
                            # Restore the saved transcriptions and inserted_transcriptions
                            st.session_state["transcriptions"] = saved_transcriptions
                            st.session_state["inserted_transcriptions"] = saved_inserted_transcriptions

                            st.session_state.pending_video_file = None
                            st.success(f'Video opened successfully! Total frames: {st.session_state.total_frames}')
                            # Set the active tab to Visual Transcription
                            st.session_state.active_tab = 2 # This might not directly switch tabs, rely on rerun and user click
                            # Store a flag to indicate we should show the workspace tab content
                            st.session_state.show_workspace = True
                            # Rerun to reflect changes and attempt tab switch
                            st.experimental_rerun()
                        else:
                            st.error('Could not open video file using OpenCV.')
                            st.session_state.uploaded = False
                            st.session_state.video = None
                            try:
                                os.unlink(temp_file_path)
                            except OSError:
                                st.warning(f"Could not delete temporary file: {temp_file_path}")


        else:
            # Display a disabled button or message when no video is uploaded
             st.write("Please upload a video before processing.") # Changed from button to text


# -----------------------------------------------
# TAB 3: VISUAL TRANSCRIPTION WORKSPACE
# -----------------------------------------------
with workspace_tab:
    # --- Display the video frame and controls only if video is uploaded and available ---
    if st.session_state.get("uploaded", False) and st.session_state.get("video") is not None:
        try:
            # --- Get Current Frame ---
            video_obj = st.session_state.video
            if video_obj is None or not video_obj.isOpened():
                st.error("Video object is not available or not opened. Please load your video in the Media Upload tab.")
                st.session_state.uploaded = False  # Reset the uploaded flag to force re-upload
            else:
                # Add prompt selection radio buttons
                st.subheader("Transcription Settings")
                st.write("Select a prompt category for image description when transcribing:")
                
                # Initialize the prompt category in session state if it doesn't exist
                if "prompt_category" not in st.session_state:
                    st.session_state.prompt_category = "general"
                
                # Create radio buttons for prompt categories
                prompt_category = st.radio(
                    "Prompt Category",
                    ["General Purpose", "STEM", "Humanities & Social Sciences", "Business"],
                    index=0,  # Default to General
                    key="prompt_radio"
                )
                
                # Update the session state based on selection
                category_mapping = {
                    "General Purpose": "general",
                    "STEM": "stem",
                    "Humanities & Social Sciences": "humanities",
                    "Business": "business"
                }
                
                # Update prompt category in session state if changed
                selected_category = category_mapping[prompt_category]
                if st.session_state.prompt_category != selected_category:
                    st.session_state.prompt_category = selected_category
                    # Load the appropriate prompt from the JSON file
                    try:
                        with open(f"database/prompts/{selected_category}.json", "r") as json_file:
                            prompt_data = json.load(json_file)
                            st.session_state['gpt-4o'] = {
                                "prompt": prompt_data["prompt"], 
                                "max_words": st.session_state["max_words"]
                            }
                            # Replace placeholder with actual max words
                            st.session_state['gpt-4o']["prompt"] = st.session_state['gpt-4o']["prompt"].replace(
                                "%MAX_WORDS%", str(st.session_state["max_words"]))
                            st.success(f"Loaded {prompt_data['name']} prompt")
                    except Exception as e:
                        st.error(f"Error loading prompt: {e}")
                
                # Display current prompt
                current_prompt = st.session_state['gpt-4o'].get("prompt", "No prompt available")
                st.markdown(f"### Current Prompt")
                st.text_area(
                    "Current prompt (resizable)",
                    value=current_prompt,
                    height=100,
                    key="prompt_display_area",
                    disabled=True,
                    label_visibility="collapsed"
                )
                
                st.markdown("---")  # Add separator
                
                video_obj.set(cv2.CAP_PROP_POS_FRAMES, st.session_state.frame_number)
                ret, frame_bgr = video_obj.read()  # Keep original frame in BGR

                if ret:
                    # --- Prepare for Canvas ---
                    # Make a copy for display conversion to avoid modifying original frame_bgr
                    frame_rgb = cv2.cvtColor(frame_bgr.copy(), cv2.COLOR_BGR2RGB)
                    pil_image_bg = Image.fromarray(frame_rgb)

                    # Initialize the crop variable to avoid NameError when checking later
                    current_processed_crop_bgr = None

                    # --- Canvas Mode and Display ---
                    use_rect_mode = st.checkbox("Use Rectangular Crop Mode", value=True, key='crop_mode_checkbox')

                    # Define canvas dimensions (use frame dimensions)
                    canvas_height, canvas_width = frame_bgr.shape[:2]
                    # Optional: Limit max display size for very large videos
                    display_width = min(canvas_width, 1080)
                    display_height = int(display_width * (canvas_height / canvas_width))  # Maintain exact aspect ratio
                    
                    st.write(f"Draw a **{'Rectangle' if use_rect_mode else 'Freeform Shape'}** on the image below.")

                    # Use the canvas_key from session state to force redraw when needed
                    current_canvas_key = f"main_canvas_video_{st.session_state.canvas_key}"
                    try:
                        canvas_result = st_canvas(
                            fill_color="rgba(255, 165, 0, 0.3)",
                            stroke_width=st.session_state.get('stroke_slider', 3), # Use stroke width from settings
                            stroke_color=st.session_state.get('stroke_color', '#00FF00'), # Use stroke color from settings
                            background_color="#eee",
                            background_image=pil_image_bg, # Use PIL image here
                            update_streamlit=True,
                            height=display_height, # Use calculated height to maintain aspect ratio
                            width=display_width,   # Use calculated display size
                            drawing_mode="rect" if use_rect_mode else "freedraw",
                            key=current_canvas_key # Dynamic key that changes to force canvas reset
                        )
                    except Exception as canvas_error:
                        st.error(f"Error rendering canvas: {canvas_error}")
                        canvas_result = None

                    # --- Process Canvas Result (no preview, just processing) ---
                    if canvas_result and canvas_result.json_data is not None and canvas_result.json_data.get("objects"):
                        last_object = canvas_result.json_data["objects"][-1]
                        # Scale factor if canvas size was different from original frame size
                        scale_x = canvas_width / display_width
                        scale_y = canvas_height / display_height

                        if use_rect_mode and last_object["type"] == "rect":
                            # Scale coordinates back to original frame dimensions
                            scaled_rect_data = {
                                'left': last_object['left'] * scale_x,
                                'top': last_object['top'] * scale_y,
                                'width': last_object['width'] * scale_x,
                                'height': last_object['height'] * scale_y,
                            }
                            current_processed_crop_bgr = crop_rectangular(frame_bgr, scaled_rect_data)

                        elif not use_rect_mode and last_object["type"] == "path":
                            # Scale path points back to original frame dimensions
                            original_path_data = []
                            for point_cmd in last_object["path"]:
                                scaled_cmd = [point_cmd[0]] # Keep command
                                # Scale coordinate values
                                for i in range(1, len(point_cmd)):
                                    scaled_cmd.append(point_cmd[i] * (scale_x if i % 2 != 0 else scale_y))
                                original_path_data.append(scaled_cmd)

                            current_processed_crop_bgr = crop_freeform(frame_bgr, original_path_data)
                            
                    # --- Frame Navigation ---
                    st.markdown("---")  # Add separator
                    fps = video_obj.get(cv2.CAP_PROP_FPS)
                    total_frames = st.session_state.total_frames
                    total_duration = total_frames / fps if fps > 0 else 0
                    
                    # Convert current frame to time
                    current_time = st.session_state.frame_number / fps if fps > 0 else 0
                    
                    # Create time-based slider
                    time_input = st.slider(
                        'Select time (seconds)',
                        0.0,
                        max(0.1, total_duration),
                        current_time,
                        step=1.0,  # Increment by 1 second
                        key='time_slider'
                    )
                    
                    # Convert selected time back to frame number
                    new_frame_number = int(time_input * fps) if fps > 0 else 0
                    
                    # Update frame number only if time selection changes
                    if new_frame_number != st.session_state.frame_number:
                        st.session_state.frame_number = new_frame_number
                        # Clear transient crop when navigating away
                        try:
                            current_processed_crop_bgr = None
                        except:
                            pass
                        # Force rerun to update the display with the new frame
                        st.experimental_rerun()

                    # Display both time and frame information
                    st.write(f"Current Time: {current_time:.2f}s (Frame: {st.session_state.frame_number})")
                    
                    # --- Navigation and Save Buttons ---
                    col1, col2, col3 = st.columns([1, 2, 1])
                    with col1:
                        if st.button(f'⬅️ Back {st.session_state.frame_increment} Frame{"s" if st.session_state.frame_increment > 1 else ""}'):
                            if st.session_state.frame_number > 0:
                                # Calculate new frame number with bounds checking
                                new_frame = st.session_state.frame_number - st.session_state.frame_increment
                                # Ensure we don't go below 0
                                st.session_state.frame_number = max(0, new_frame)
                                # Clear transient crop when navigating away
                                try:
                                    current_processed_crop_bgr = None
                                except:
                                    pass
                                st.experimental_rerun()

                    with col2:
                        # Save Button
                        if st.button('💾 Save Frame (Crop if Drawn)'):
                            # Re-get the frame to ensure it's the one displayed
                            video_obj = st.session_state.video
                            if video_obj and video_obj.isOpened():
                                video_obj.set(cv2.CAP_PROP_POS_FRAMES, st.session_state.frame_number)
                                ret_save, frame_bgr_save = video_obj.read()

                                if ret_save:
                                    saved_image_data_rgb = None
                                    is_cropped_flag = False

                                    # Check if there's a crop to use
                                    if current_processed_crop_bgr is not None and current_processed_crop_bgr.size > 0:
                                        # Save the cropped version (convert to RGB)
                                        saved_image_data_rgb = cv2.cvtColor(current_processed_crop_bgr, cv2.COLOR_BGR2RGB)
                                        is_cropped_flag = True
                                        st.success(f"Saving **cropped** frame {st.session_state.frame_number}")
                                    else:
                                        # Save the full frame version (convert to RGB)
                                        saved_image_data_rgb = cv2.cvtColor(frame_bgr_save, cv2.COLOR_BGR2RGB)
                                        is_cropped_flag = False
                                        st.success(f"Saving **full** frame {st.session_state.frame_number}")

                                    # Store in session state with frame info
                                    try:
                                        # Create a copy of the numpy array
                                        frame_copy = saved_image_data_rgb.copy()
                                        
                                        # Store in session state with clear structure
                                        st.session_state.saved_frames[st.session_state.frame_number] = {
                                            'frame': frame_copy, # Store RGB numpy array copy
                                            'frame_number': st.session_state.frame_number,
                                            'is_cropped': is_cropped_flag,
                                            'has_visual_transcripts': False,
                                            'getting_visual_transcripts': False,
                                            'visual_transcripts': None
                                        }
                                        
                                        # Increment the canvas key to force a redraw/clear of the canvas
                                        st.session_state.canvas_key += 1
                                        
                                        # Clear the crop preview after saving
                                        current_processed_crop_bgr = None
                                        
                                        st.success(f"Saved frame {st.session_state.frame_number}")
                                        st.experimental_rerun()
                                    except Exception as save_error:
                                        st.error(f"Error saving frame: {save_error}")
                                else:
                                    st.error('Could not capture the frame to save.')
                            else:
                                st.error("Video is not available. Please load your video in the Media Upload tab.")

                    with col3:
                        if st.button(f'➡️ Forward {st.session_state.frame_increment} Frame{"s" if st.session_state.frame_increment > 1 else ""}'):
                            if st.session_state.frame_number < st.session_state.total_frames - 1:
                                # Calculate new frame number with bounds checking
                                new_frame = st.session_state.frame_number + st.session_state.frame_increment
                                # Ensure we don't go beyond the last frame
                                st.session_state.frame_number = min(st.session_state.total_frames - 1, new_frame)
                                # Clear transient crop when navigating away
                                try:
                                    current_processed_crop_bgr = None
                                except:
                                    pass
                                st.experimental_rerun()
                else:
                    st.error(f'Could not read frame {st.session_state.frame_number}. End of video or error.')
                    
                    # Add a button to allow resetting the video
                    if st.button("Reset Video"):
                        if "video" in st.session_state:
                            try:
                                st.session_state.video.release()
                            except:
                                pass
                        st.session_state.video = None
                        st.session_state.uploaded = False
                        st.experimental_rerun()
        except Exception as e:
            st.error(f"Error accessing or processing video: {e}")
            st.info("Please try uploading your video again in the Media Upload tab.")
            # Reset video-related session state to force clean reload
            st.session_state.uploaded = False
            if "video" in st.session_state:
                try:
                    st.session_state.video.release()
                except:
                    pass
                st.session_state.video = None
    else:
        # Display a placeholder when no video is uploaded
        try:
            placeholder_image_path = "\image_place_holder.png"
            if os.path.exists(placeholder_image_path):
                # Display the placeholder image
                st.image(placeholder_image_path, use_column_width=True, caption="Please upload a video in the Media Upload tab to begin.")
            else:
                st.info("Please upload a video in the Media Upload tab to begin.")
        except Exception as e:
            st.info("Please upload a video in the Media Upload tab to begin.")
            
        # Add a button to allow resetting the session state if needed
        if st.session_state.get("uploaded", False) or st.session_state.get("video") is not None:
            if st.button("Reset Video State"):
                # Clean up video if it exists
                if "video" in st.session_state and st.session_state.video is not None:
                    try:
                        st.session_state.video.release()
                    except:
                        pass
                # Reset all video-related state variables
                st.session_state.video = None
                st.session_state.uploaded = False
                st.session_state.frame_number = 0
                st.experimental_rerun()
    
    # --- Display Final Transcript Information (moved into the workspace tab) ---
    st.markdown("---")
    st.subheader("Combined Transcript (Preview)")
    
    # Use our new merge_transcripts function to get combined chronological transcript
    if st.session_state.audio_transcript:
        st.write(st.session_state.audio_transcript)
    else:
        # Get merged transcripts in chronological order
        merged_transcripts = merge_transcripts()
        
        if merged_transcripts:
            st.write("Showing all transcripts in chronological order:")
            
            # Display merged transcripts
            for entry in merged_transcripts:
                if len(entry) == 3:  # Audio entry (timestamp, text, "audio")
                    timestamp, text, source = entry
                    st.write(f"**[{timestamp:.2f}s]** {text}")
                else:  # Visual entry (timestamp, text, "visual", frame_number)
                    timestamp, text, source, frame_number = entry
                    st.write(f"**[{timestamp:.2f}s - Frame {frame_number}]** 🖼️ *{text}*")
                
            # Add a download option for the combined transcript
            if st.button("Update Combined Transcript Document"):
                # Create a new document with merged transcripts
                doc = Document()
                doc.add_heading("Combined Visual and Audio Transcript", level=1)
                for entry in merged_transcripts:
                    if len(entry) == 3:  # Audio entry
                        timestamp, text, source = entry
                        doc.add_paragraph(f"[{timestamp:.2f}s] {text}")
                    else:  # Visual entry
                        timestamp, text, source, frame_number = entry
                        doc.add_paragraph(f"[{timestamp:.2f}s - Frame {frame_number}] {text}")
                
                # Save and offer for download
                temp_doc_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
                doc.save(temp_doc_path)
                with open(temp_doc_path, "rb") as doc_file:
                    st.download_button(
                        "Download Combined Transcript", 
                        doc_file, 
                        file_name="combined_transcript.docx", 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else:
            st.info("No transcript data available. Please upload an SRT file in the Media Upload tab or add frame transcriptions.")

# -----------------------------------------------
# Sidebar: Display Saved Frames 
# -----------------------------------------------
with st.sidebar:
    st.markdown("### Selected Frames for Transcription")
    
    # Check if saved_frames exists and is not empty
    if not st.session_state.get("saved_frames", {}):
         st.write("No frames selected yet.")
    else:
        # Sort frames for consistent display by frame number
        for frame_number in sorted(st.session_state.saved_frames.keys()):
            frame_info = st.session_state.saved_frames[frame_number]
            
            # Use helper function to get base64
            try:
                 # Re-create base64 string each time to avoid caching issues
                 if 'frame' in frame_info and frame_info['frame'] is not None:
                     base64_img = image_to_base64(frame_info['frame']) # Assumes frame is RGB numpy array
                 else:
                     st.warning(f"Frame {frame_number} data is missing")
                     # Skip this frame to avoid display errors
                     continue
            except Exception as e:
                 st.error(f"Error converting Frame {frame_number} to base64: {e}")
                 # Skip this frame if conversion fails
                 continue

            transcript_text = frame_info.get('visual_transcripts', 'No transcript yet') # Use .get for safety
            is_cropped_text = "(Cropped)" if frame_info.get('is_cropped', False) else "(Full Frame)" # Check crop status

            # Display Card using Markdown with error handling
            try:
                st.markdown(f"""
                    <div style="border: 1px solid #ccc; padding: 10px; margin-bottom: 10px; border-radius: 5px;">
                        <h5>Frame {frame_number} {is_cropped_text}</h5>
                        <img src="data:image/jpeg;base64,{base64_img}" style="width:100%;" alt="Frame {frame_number}" />
                        <p style="font-size: smaller; white-space: pre-wrap; word-wrap: break-word;">{transcript_text}</p>
                    </div>
                """, unsafe_allow_html=True)
            except Exception as display_error:
                st.error(f"Error displaying Frame {frame_number}: {display_error}")
                continue

            # Buttons for actions
            col1_side, col2_side = st.columns(2)
            with col1_side:
                # Transcribe Button (only if not already transcribed)
                if not frame_info.get('has_visual_transcripts', False):
                    if st.button(f"Transcribe #{frame_number}", key=f"btn_{frame_number}"):
                        st.info(f"Transcribing Frame {frame_number}...")
                        try:
                            # Pass the saved frame data (numpy array) directly
                            saved_frame_data = frame_info['frame'] # This is already RGB
                            
                            # Get the image as base64 for OpenAI API
                            img_pil = Image.fromarray(saved_frame_data)
                            base64_image = encode_image(img_pil)
                            
                            # Use the GPT-4o prompt from settings
                            prompt_text = st.session_state['gpt-4o'].get("prompt", "What's in this image?")
                            prompt_category = st.session_state.get("prompt_category", "general")
                            
                            # Display which prompt is being used
                            category_name_map = {
                                "general": "General Purpose",
                                "stem": "STEM",
                                "humanities": "Humanities & Social Sciences",
                                "business": "Business"
                            }
                            category_name = category_name_map.get(prompt_category, "Custom")
                            st.info(f"Using {category_name} prompt: \"{prompt_text}\"")
                            
                            # Show max tokens information
                            max_tokens = int(st.session_state["max_words"]) * 4
                            st.info(f"Using word limit from configuration: {st.session_state['max_words']} tokens")
                            
                            # Prepare headers and payload for OpenAI API
                            headers = {"Content-Type": "application/json", "Authorization": f"Bearer {GPT_API_KEY}"}
                            payload = {
                                "model": "gpt-4o",
                                "messages": [
                                    {"role": "user", "content": [
                                        {"type": "text", "text": prompt_text},
                                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                                    ]}
                                ],
                                "max_tokens": int(st.session_state["max_words"]) * 4  # Convert word count to approximate token count (4 tokens per word on average)
                            }
                            
                            # Make API call
                            response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
                            gpt_response = response.json()
                            transcription = gpt_response['choices'][0]['message']['content']

                            # Update session state
                            st.session_state.saved_frames[frame_number]['visual_transcripts'] = transcription
                            st.session_state.saved_frames[frame_number]['has_visual_transcripts'] = True
                            st.session_state["transcriptions"][frame_number] = transcription
                            
                            # Get timestamp if needed
                            if st.session_state.get('video') and st.session_state.get('video').isOpened():
                                try:
                                    video_obj = st.session_state.video
                                    st.session_state.saved_frames[frame_number]['time_stamp'] = get_frame_timestamp(frame_number, video_obj)
                                except Exception as ts_error:
                                    st.warning(f"Could not get timestamp for Frame {frame_number}: {ts_error}")
                                    st.session_state.saved_frames[frame_number]['time_stamp'] = "N/A"

                            st.success(f"Transcription completed for Frame {frame_number}.")
                            st.experimental_rerun()  # Update sidebar display
                        except Exception as api_error:
                            st.error(f"Transcription failed: {api_error}")

            with col2_side:
                # Insert into Transcript / Remove Button
                if frame_info.get('has_visual_transcripts', False):
                    if st.button(f"Insert to Transcript #{frame_number}", key=f"add_{frame_number}"):
                        try:
                            # When a visual transcript is inserted:
                            # 1. It's added to the existing subtitle text (if matching frame exists)
                            # 2. The frame number is added to inserted_transcriptions set
                            # 3. The merge_transcripts() function will include it in the combined chronological display
                            
                            # Insert into the subtitles dictionary
                            if frame_number in st.session_state["frame_subtitle_map"]:
                                # Get existing subtitle text
                                subtitle_key = None
                                for start_time, text in st.session_state["subtitles"].items():
                                    if int(start_time * int(st.session_state.video.get(cv2.CAP_PROP_FPS))) == frame_number:
                                        subtitle_key = start_time
                                        break
                                
                                if subtitle_key is not None:
                                    # Add the GPT transcription to the subtitle
                                    st.session_state["subtitles"][subtitle_key] += f"\n[GPT]: {st.session_state['transcriptions'][frame_number]}"
                                    # Add this frame to the set of inserted transcriptions
                                    st.session_state.inserted_transcriptions.add(frame_number)
                                    st.success(f"Inserted GPT transcription into frame {frame_number} subtitle.")
                                else:
                                    st.warning(f"Could not find subtitle for frame {frame_number}.")
                            else:
                                # Even without a subtitle mapping, we can still track that this frame's 
                                # transcription has been inserted
                                st.session_state.inserted_transcriptions.add(frame_number)
                                st.warning(f"No subtitle mapping found for frame {frame_number}, but marked as inserted.")
                            
                            # Try to use the insert_VT_into_AT utility if available
                            try:
                                if 'insert_VT_into_AT' in globals() and st.session_state.audio_transcript:
                                    insert_VT_into_AT(frame_info)
                                    st.success(f"Also added Frame {frame_number} info to audio transcript.")
                            except Exception as insert_err:
                                st.warning(f"Could not insert into audio transcript: {insert_err}")
                                
                            st.experimental_rerun()
                        except Exception as insert_err:
                            st.error(f"Error inserting transcription: {insert_err}")

                # Add a button to remove a frame from selection
                if st.button(f"Remove #{frame_number}", key=f"del_{frame_number}"):
                    if frame_number in st.session_state.saved_frames:
                        # Remove the frame from saved_frames
                        del st.session_state.saved_frames[frame_number]
                        
                        # Remove from inserted_transcriptions if it was inserted
                        if frame_number in st.session_state.inserted_transcriptions:
                            st.session_state.inserted_transcriptions.remove(frame_number)
                        
                        # Remove from transcriptions dictionary
                        if frame_number in st.session_state.transcriptions:
                            del st.session_state.transcriptions[frame_number]
                            
                        # If this frame has a corresponding subtitle entry, clean up the GPT part
                        if frame_number in st.session_state["frame_subtitle_map"]:
                            # Find the subtitle key
                            subtitle_key = None
                            for start_time, text in st.session_state["subtitles"].items():
                                if int(start_time * int(st.session_state.video.get(cv2.CAP_PROP_FPS))) == frame_number:
                                    subtitle_key = start_time
                                    break
                            
                            if subtitle_key is not None and "\n[GPT]:" in st.session_state["subtitles"][subtitle_key]:
                                # Remove only the GPT part
                                original_text = st.session_state["subtitles"][subtitle_key].split("\n[GPT]:")[0]
                                st.session_state["subtitles"][subtitle_key] = original_text
                        
                        st.success(f"Removed Frame {frame_number} and its transcription.")
                        st.experimental_rerun()  # Update sidebar

    # Download options
    st.sidebar.subheader("Download Options")
    download_transcript()

# Add spacing at the bottom
st.markdown("<div style='height: 100px;'></div>", unsafe_allow_html=True)
